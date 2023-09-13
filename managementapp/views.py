from email import message_from_binary_file
from email.headerregistry import Group
from django.conf import settings
from django.http import HttpResponse, HttpResponseRedirect, request, Http404, JsonResponse, HttpResponseBadRequest
from django.shortcuts import render, redirect, get_object_or_404
from django.views.generic import ListView, DetailView, CreateView, TemplateView, View
from django.core.mail import EmailMessage, send_mass_mail, send_mail, BadHeaderError
from django.contrib.auth.decorators import login_required, user_passes_test
from django.core.exceptions import ObjectDoesNotExist
from django.contrib import messages
from django.urls import reverse_lazy, reverse
from django.template.defaulttags import register
from datetime import date, datetime, timedelta, timezone
from django.contrib.auth.base_user import BaseUserManager
from django.db.models import Sum, Q
from django.contrib.auth import logout
from django.contrib.auth.forms import AuthenticationForm
from django.contrib.auth.views import LoginView
from django.contrib.auth import login, authenticate
from openpyxl import load_workbook, Workbook
from openpyxl.drawing.image import Image
from openpyxl.styles import Alignment
from openpyxl.utils import get_column_letter
from django.core.files.storage import FileSystemStorage
from django.contrib.auth.mixins import UserPassesTestMixin
from cloitLMS.settings import BASE_DIR
from .models import *
from examapp.models import *
from learningapp.models import *
from learningapp.views import get_learning_time, get_studied_questions, question_count, get_percent
from django.utils.formats import localize
import json
import requests
import re
import time
import docx

symbols = (u"abcdefghijklmnoprstuvwxyzABCDEFGHIJKLMNOPRSTUVWXYZ.-_@1234567890",
        u"123456789012345678901234512345678901234567890123450000ABCDEFGHIJ")
tr = {ord(a):ord(b) for a, b in zip(*symbols)} 

def user_message(login, password, first_name, middle_name, path):
    message = (
        'Уважаемый ' + first_name + ' ' + middle_name + '!\n\n' +
        'Вам предоставлен доступ к сервису обучения и тренинга «Школа главного инженера»\n\n' +
        
        'Для доступа к Вашему личному кабинету Вам необходимо перейти по ссылке https://' + path + '/login\n' +
        'Логин: ' + login + '\n' +
        'Пароль: ' + password + '\n\n' +
        'Инструкция пользователя доступна по ссылке: https://' + path + '/media/Инструкция%20обучающегося.pdf\n' +
        'В случае если у Вас возникнут вопросы - Вы можете обратиться в техническую поддержку:\n' +
        '- Email: support@sdo-vot.ru\n' +
        '- Tel: +7(495) 320-88-07 ( с 06:00 по 23:00 по Московскому времени)\n' +

        'Обращаем Ваше внимание, что логин и пароль вы можете скопировать из данного письма и вставить в окне авторизации.'
        )
    print(message)
    return message

@register.filter
def qty_sum(user):
    qty = Subscription.objects.filter(manager = user).aggregate(Sum('qty'))
    return 0 if qty['qty__sum'] is None else qty['qty__sum']

@register.filter
def check_attempt_exist(value, user):
    return LearningQuizAttempt.objects.filter(user = user, quiz = Quiz.objects.get(course = value, name__contains = 'Проверка знаний'), status_id = 3).exists()

@register.filter
def last_login(user):
    return User.objects.get(pk = user).last_login

@register.filter
def check_home_page(user):
    if user.groups.filter(name='manager').exists():
        return str('/users')
    elif user.is_staff:
        return str('/adm/managerlist')
    else:
        return str('/learning')

@register.filter
def split_parent_course(value, course):
    return value.replace(course, '')

@register.filter
def qty_month_sum(user):
    today = datetime.datetime.today()
    datem = datetime.datetime(today.year, today.month, 1)
    qty = Subscription.objects.filter(manager = user, date__gte = datem, qty__lt = 0).count()
    return 0 if qty is None else qty

@register.filter
def check_kaiten_messages(user):
    headers = {
        'Accept': 'application/json', 
        'Content-Type': 'application/json', 
        'Authorization': 'Bearer 680c3c04-c773-48b7-ba1f-9a5c9951ca95'
    }
    cards = HelpDeskCards.objects.filter(user = user)
    for card in reversed(cards):
        url = f"https://alexeyomich.kaiten.ru/api/latest/cards/{card.card}/comments"
        r = requests.get(url, headers=headers)
        if r.status_code == 429:
            time.sleep(1)
            r = requests.get(url, headers=headers)
        if r.status_code == 200:
            if r.json() and r.json()[-1]['author']['full_name'] != 'Guest' and r.json()[-1]['created'] > card.update:

                return True
    
    return False

def About(request):
    return render(request, 'about.html')

def logout_user(request):
    logout(request)
    return redirect('managementapp:login')

class CheckManagerPermissions(UserPassesTestMixin):
    
    def test_func(self):
        return self.request.user.groups.filter(name = 'manager').exists()
    
    def handle_no_permission(self):
        return redirect('learningapp:courselist')

class CheckSalerPermissions(UserPassesTestMixin):
    
    def test_func(self):
        return self.request.user.groups.filter(name = 'saler').exists()
    
    def handle_no_permission(self):
        return redirect('learningapp:courselist')
    
class Home(TemplateView):
    template_name = 'home.html'
        
    def get_context_data(self, *args, **kwargs):
        context = super().get_context_data(**kwargs)
        
        return context

class Support(TemplateView):
    template_name = 'support.html'

    def post(self, request, *args, **kwargs):
        
        if request.POST.get("form_type") == 'send_email':

            description = f"ФИО: {request.POST.get('name')}\n\nЛогин: {request.POST.get('login')}\n\nEmail: {request.POST.get('email')}\n\nСообщение: {request.POST.get('text')}"

            url = "https://alexeyomich.kaiten.ru/api/latest/cards"
            
            data = { 
                "title": "Заявка с Формы ШГИ", 
                "board_id": 468128, 
                "description": description,
                "responsible_id": 417435, 
                "external_id": (self.request.user.id if self.request.user.is_authenticated else ''), 
                "properties": {}
            }
        
            headers = {
                'Accept': 'application/json', 
                'Content-Type': 'application/json', 
                'Authorization': 'Bearer 9f4ff229-aee5-43f2-a765-18e93a818c7e'
            }
            
            r = requests.post(url, data=json.dumps(data), headers=headers)
            if self.request.user.is_authenticated:
                try:
                    HelpDeskCards.objects.create(user = self.request.user, card = int(r.json()['id']), update = r.json()['created'])

                    return redirect('managementapp:supportcards')
                
                except:
                    pass

        return redirect('managementapp:home')

@register.filter
def card_get_id(card):
    return card.split(' ')[1]

def support_card_detail(card, update):
    url = f"https://alexeyomich.kaiten.ru/api/latest/cards/{card}"
    headers = {
        'Accept': 'application/json', 
        'Content-Type': 'application/json', 
        'Authorization': 'Bearer 680c3c04-c773-48b7-ba1f-9a5c9951ca95'
    }
    r = requests.get(url, headers=headers)
    if r.status_code == 429:
        time.sleep(1)
        r = requests.get(url, headers=headers)
    if r.json()['comment_last_added_at'] and r.json()['comment_last_added_at'] > update:
        text = f"Заявка: {r.json()['id']} от {datetime.datetime.strftime(datetime.datetime.fromisoformat(r.json()['created']), '%d.%m.%Y %H:%M')}; Статус: {r.json()['column']['title']}<br>{(r.json()['description']).split('Сообщение: ')[1]}<br><br>Непрочитанное сообщение!"
    else:
        text = f"Заявка: {r.json()['id']} от {datetime.datetime.strftime(datetime.datetime.fromisoformat(r.json()['created']), '%d.%m.%Y %H:%M')}; Статус: {r.json()['column']['title']}<br>{(r.json()['description']).split('Сообщение: ')[1]}<br>"
    
    return text

def support_card_comments(card, user):
    url = f"https://alexeyomich.kaiten.ru/api/latest/cards/{card}/comments"
    headers = {
        'Accept': 'application/json', 
        'Content-Type': 'application/json', 
        'Authorization': 'Bearer 680c3c04-c773-48b7-ba1f-9a5c9951ca95'
    }
    r = requests.get(url, headers=headers)
    comments = []
    comments_dict = {}
    for comment in r.json():
        if comment['author']['full_name'] == 'Guest':
            comments_dict[comment['id']] = f"{user}:<br>{comment['text']}"
        else:
            comments_dict[comment['id']] = f"{comment['author']['full_name']}:<br>{comment['text']}"
    for key in sorted(comments_dict):
        comments.append(comments_dict[key])
    hdc = HelpDeskCards.objects.get(card = card, user = user)
    try:
        if hdc.update < r.json()[-1]['created']:
            hdc.update = r.json()[-1]['created']
            hdc.save()
    except:
        pass
    return comments

class SupportCards(ListView):
    model = HelpDeskCards
    template_name = 'support cards.html'

    def get_queryset(self, *args, **kwargs):
        if self.request.GET.get('card'):
            return HelpDeskCards.objects.get(card = self.request.GET.get('card'))
        else:
            cards = HelpDeskCards.objects.filter(user = self.request.user)
            cards_list = []
            for card in cards:
                cards_list.append(support_card_detail(card.card, card.update))
            return cards_list
    
    def post(self, request, *args, **kwargs):
        if request.POST.get("form_type") == 'send' and request.POST.get('text'):
            url = f"https://alexeyomich.kaiten.ru/api/latest/cards/{self.request.GET.get('card')}/comments"
            headers = {
                'Accept': 'application/json', 
                'Content-Type': 'application/json', 
                'Authorization': 'Bearer 680c3c04-c773-48b7-ba1f-9a5c9951ca95'
            }
            data = { 
                "text": request.POST.get('text'), 
            }            
            r = requests.post(url, data=json.dumps(data), headers=headers)
            hdc = HelpDeskCards.objects.get(card = self.request.GET.get('card'))
            hdc.update = r.json()['updated']
            hdc.save()

        return HttpResponseRedirect("?card=%s" % self.request.GET.get('card'))

    def get_context_data(self, *args, **kwargs):
        context = super().get_context_data(**kwargs)
        if self.request.GET.get('card'):
            context['card'] = self.request.GET.get('card')
            context['comments'] = support_card_comments(self.request.GET.get('card'), self.request.user)
        
        return context

class LogInView(LoginView):
    form_class = AuthenticationForm
    template_name = 'login.html'

    def get_success_url(self):
        return reverse_lazy('managementapp:users')

    def post(self, request, **kwargs):

        user = authenticate(username = request.POST.get('login'), password = request.POST.get('password'))
        if user is not None:
            if user.is_active:
                login(request, user)
                if self.request.user.groups.filter(name='manager').exists():
                    return redirect('managementapp:users')
                if self.request.user.groups.filter(name='saler').exists():
                    return redirect('managementapp:managers')
                elif self.request.user.is_staff:
                    return redirect('administratorapp:managerlist')
                else:
                    return redirect('learningapp:courselist')
            else:
                messages.warning(request, 'Учетная запись заблокирована!')
                return redirect('managementapp:login')
        else:
            messages.error(request, 'Пользователь или пароль не верны!')
            return redirect('managementapp:login')
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)

        return context
    
class UsersRegistration(CheckManagerPermissions, TemplateView):
    template_name = 'managementapp/registration.html'

    def post(self, request, *args, **kwargs):

        ''' reCAPTCHA validation '''
        recaptcha_response = request.POST.get('g-recaptcha-response')
        data = {
        'secret': '6LfkYW0eAAAAANaVAWFKj83gjWJhvwe-OECy3MDu',
        'response': recaptcha_response
        }
        r = requests.post('https://www.google.com/recaptcha/api/siteverify', data=data)

        password = BaseUserManager().make_random_password(length=5, allowed_chars='abcdefghjkmnpqrstuvwxyz0123456789')
        password = password + BaseUserManager().make_random_password(length=5, allowed_chars='ABCDEFGHJKLMNPQRSTUVWXYZ0123456789')
        
        last_group, create = Group.objects.get_or_create(
            name = self.request.user
        )

        last_user, create = User.objects.get_or_create(
            username = (str(self.request.user).split('@')[0] + '_user' + str(User.objects.filter(pk__in = last_group.user_set.all()).count())),
            first_name = request.POST.get('firstname'),
            last_name = request.POST.get('lastname'),
            email = request.POST.get('email'),
            is_superuser = False,
            is_staff = False,
            is_active = True,
        )
        if create:
            last_user.set_password(password)  
            last_user.save()

        last_group.user_set.add(last_user)

        extend_user, create = ExtendingUserFields.objects.get_or_create(
            user = last_user,
            middle_name = request.POST.get('middlename'),
            department = request.POST.get('department'),
            password = password,
        )

        return redirect('managementapp:users')
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)

        return context

class EditUser(CheckManagerPermissions, DetailView):
    model = User
    context_object_name = 'object'
    template_name = 'managementapp/edituser.html'

    def get_object(self, **kwargs):
        
        return ExtendingUserFields.objects.get(user = self.kwargs['user_id'])

    def post(self, request, *args, **kwargs):
        update_user = User.objects.get(pk = self.kwargs['user_id'])
            
        ''' reCAPTCHA validation '''
        recaptcha_response = request.POST.get('g-recaptcha-response')
        data = {
        'secret': '6LfkYW0eAAAAANaVAWFKj83gjWJhvwe-OECy3MDu',
        'response': recaptcha_response
        }
        r = requests.post('https://www.google.com/recaptcha/api/siteverify', data=data)

        update_user = User.objects.get(pk = self.kwargs['user_id'])
        update_user.first_name = request.POST.get('firstname')
        update_user.last_name = request.POST.get('lastname')
        update_user.email = request.POST.get('email')
        update_user.save()

        update_extend_user = ExtendingUserFields.objects.get(user = self.kwargs['user_id'])
        update_extend_user.middle_name = request.POST.get('middlename')
        update_extend_user.department = request.POST.get('department')
        update_extend_user.save()

        return redirect('managementapp:users')

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)

        return context

class ManagersList(CheckSalerPermissions, ListView):
    model = User
    paginate_by = 10
    template_name = 'managementapp/managerlist.html'
    context_object_name = 'object_list'
    
    def get_queryset(self, **kwargs):
        query = self.request.GET.get('search')
        if query == None:
            obj = (
                User.objects
                .filter(groups__name = 'manager').filter(groups__name = self.request.user)
                .select_related('extendinguserfields')
                .order_by('-id')
                )
            return obj
        else:
            obj = (
                User.objects
                .filter(groups__name = 'manager').filter(groups__name = self.request.user)
                .filter(Q(username__icontains = query) | Q(first_name__icontains = query) | Q(last_name__icontains = query) | Q(email__icontains = query) | Q(extendinguserfields__middle_name__icontains = query) | Q(extendinguserfields__department__icontains = query))
                .select_related('extendinguserfields')
                .order_by('-id')
                )
            self.paginate_by = obj.count()
            return obj                

    def post(self, request, *args, **kwargs):
        
        if request.POST.get("form_type") == 'add_manager' and request.POST.get('username') != None and request.POST.get('first_name') != None:
            
            if User.objects.filter(username = request.POST.get('username')).exists():
                messages.warning(request, 'Пользователь с таким логином уже существует!')
                return redirect('managementapp:managers')
        
            manager_groupe = Group.objects.get(name = 'manager')
            saler_groupe, create = Group.objects.get_or_create(name = self.request.user)            
            
            password = BaseUserManager().make_random_password(length=5, allowed_chars='abcdefghjkmnpqrstuvwxyz0123456789')
            password = password + BaseUserManager().make_random_password(length=5, allowed_chars='ABCDEFGHJKLMNPQRSTUVWXYZ0123456789')
            
            last_user, create = User.objects.get_or_create(
                username = str(request.POST.get('username')),
                first_name = str(request.POST.get('first_name')),
                last_name = '',
                email = request.POST.get('email'),
                is_superuser = False,
                is_staff = False,
                is_active = True,
            )
            
            if create:
                last_user.set_password(password)  
                last_user.save()

            manager_groupe.user_set.add(last_user)
            saler_groupe.user_set.add(last_user)

            extend_user, create = ExtendingUserFields.objects.get_or_create(
                user = last_user,
                middle_name = '',
                department = str(request.POST.get('first_name')),
                password = password,
            )
            
            return redirect('managementapp:managers')
        
        if request.POST.get("form_type") == 'add_subs' and request.POST.getlist('checkeduser') != []:
            
            if len(request.POST.getlist('checkeduser')) == 1:
                
                checked_user = User.objects.get(pk = request.POST.getlist('checkeduser')[0])
                
                if request.POST.get('trening'):
                    Subscription.objects.create(qty = int(request.POST.get('trening')), subtype = 'Тренинг ВОТ', manager = checked_user)
                                
                return redirect('managementapp:managers')
            
            elif len(request.POST.getlist('checkeduser')) > 1:
                messages.warning(request, 'Выбирите 1 пользователя!')
                return redirect('managementapp:managers')
        
        if request.POST.get("form_type") == 'report':

            start = request.POST.get('start') + ' 00:00:01'
            end = request.POST.get('end') + ' 23:59:59'

            if end >= start:
                
                checked_users = request.POST.getlist('checkeduser')
                check_all = request.POST.getlist('check_all')

                if checked_users:
                    checked_managers = User.objects.filter(pk__in=checked_users)
                    subs = Subscription.objects.filter(
                        manager__in = checked_managers, 
                        date__range = (start, end)
                    ).select_related('user', 'course', 'user__extendinguserfields')
                elif check_all:
                    subs = Subscription.objects.filter(
                        manager__in = User.objects.filter(groups__name = self.request.user),
                        date__range = (start, end)
                    ).select_related('user', 'course', 'user__extendinguserfields')
                else:
                    return HttpResponseBadRequest()

                wb = Workbook()
                ws = wb.active
                ws.append(['Логин', 'Организация', 'Подписка', 'Кол-во', 'Пользователь', 'Курс', 'Дата'])
                for sub in subs:
                    ws.append([
                        sub.manager.username, 
                        sub.manager.first_name, 
                        sub.subtype, 
                        sub.qty, 
                        (sub.user.get_full_name() if sub.user is not None else ''), 
                        (sub.course.name if sub.course is not None else ''), 
                        sub.date
                    ])
                
                response = HttpResponse(content_type='application/vnd.ms-excel')
                response['Content-Disposition'] = 'attachment;filename="report.xlsx"'
                wb.save(response)

                return response
            
            return redirect('managementapp:managers')
        
        if request.POST.get("form_type") == 'search_btn' and request.POST.get('search') != []:
            
            return HttpResponseRedirect("?search=%s" % request.POST.get('search'))
    
    def get_context_data(self, *args, **kwargs):
        context = super().get_context_data(**kwargs)

        return context

class UsersList(CheckManagerPermissions, ListView):
    model = User
    template_name = 'managementapp/users.html'
    paginate_by = 10
    context_object_name = 'object_list'
    
    def get_queryset(self, **kwargs):
        curr_group, create = Group.objects.get_or_create(name = self.request.user)
        try:
            query = self.request.GET.get('search')
            if query == None:
                obj = (
                    User.objects
                    .filter(pk__in = curr_group.user_set.all())
                    .select_related('extendinguserfields')
                    .order_by('-id')
                    )
                return obj
            else:
                obj = (
                    User.objects
                    .filter(pk__in = curr_group.user_set.all())
                    .filter(Q(username__icontains = query) | Q(first_name__icontains = query) | Q(last_name__icontains = query) | Q(email__icontains = query) | Q(extendinguserfields__middle_name__icontains = query) | Q(extendinguserfields__department__icontains = query))
                    .select_related('extendinguserfields')
                    .order_by('-id')
                    )
                self.paginate_by = obj.count()
                return obj                
        except:
            return []

    def post(self, request, *args, **kwargs):
        
        if request.POST.get("form_type") == 'edituser' and request.POST.getlist('checkeduser') != []:
            
            if len(request.POST.getlist('checkeduser')) == 1:
                return redirect('managementapp:edituser', user_id = request.POST.getlist('checkeduser')[0])
            
            elif len(request.POST.getlist('checkeduser')) > 1:
                messages.warning(request, 'Выбирите 1 пользователя!')
                return redirect('managementapp:users')
       
        if request.POST.get("form_type") == 'report' and 'formreport':

            start = request.POST.get('start') + ' 00:00:01'
            end = request.POST.get('end') + ' 23:59:59'

            if end >= start:
                
                subs = Subscription.objects.filter(
                    manager = self.request.user,
                    date__range=(start, end),
                    qty__lt = 0,
                ).select_related('user', 'course', 'user__extendinguserfields')

                wb = Workbook()
                ws = wb.active
                ws.append(['Дата', 'Курс', 'ФИО', 'Логин'])
                for sub in subs:
                    ws.append([
                        sub.date,
                        (sub.course.name if sub.course is not None else ''),
                        (sub.user.get_full_name() if sub.user is not None else ''),
                        (sub.user.username if sub.user is not None else ''),
                    ])
                ws.append(['', '', '', ''])
                ws.append([f'Итого списано подписок: {subs.count()}'])
                response = HttpResponse(content_type='application/vnd.ms-excel')
                response['Content-Disposition'] = 'attachment;filename="report.xlsx"'
                wb.save(response)

                return response
            
            else:

                messages.warning(request, 'Дата начала раньше даты конца периода!')
                return redirect('managementapp:users')
            
        if request.POST.get("form_type") == 'formreportgroup' and request.POST.getlist('checkeduser') != []:

            doc = docx.Document(str(BASE_DIR) + '/media/report-example.docx')
            
            styles = doc.styles
            styles.add_style('Heading 1', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH, builtin=True)
            
            p = doc.add_paragraph()
            runner = p.add_run('Раздел I. Перечень обучающихся и областей подготовки:')
            runner.bold = True
            
            items = []
            users = User.objects.filter(id__in = request.POST.getlist('checkeduser'))
            for idx_user, user in enumerate(users, start=1):
                courses = UserRegisteredCourse.objects.filter(user = user).select_related('course').order_by('-timeregistered').values_list('course__name', flat = True)
                
                courses_list = []
                for idx_course, course in enumerate(courses, start=1):
                    courses_list.append(f"{idx_course}. {course}")

                items.append([idx_user, f"{user.last_name} {user.first_name} {ExtendingUserFields.objects.get(user = user).middle_name}", '\n\n'.join(courses_list)])

            table = doc.add_table(1, len(items[0]))
            table.style = 'Table Grid'
            head_cells = table.rows[0].cells
            for i, item in enumerate(['№', 'ФИО', 'Курс подготовки']):
                p = head_cells[i].paragraphs[0]
                p.add_run(item).bold = True
                p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            for row in items:
                cells = table.add_row().cells
                for i, item in enumerate(row):
                    cells[i].text = str(item)
                    if i == 2:
                        cells[i].paragraphs[0].runs[0].font.name = 'Arial'
            
            for cell in table.columns[0].cells:
                cell.width = docx.shared.Cm(1)
            
            for cell in table.columns[1].cells:
                cell.width = docx.shared.Cm(5)
            
            for cell in table.columns[2].cells:
                cell.width = docx.shared.Cm(11)
            
            doc.add_page_break()
            
            p = doc.add_paragraph()
            runner = p.add_run('Раздел II. Итоги подготовки:')
            runner.bold = True

            for idx_user, user in enumerate(users, start=1):
                doc.add_paragraph()
                p = doc.add_paragraph()
                runner = p.add_run((f"{idx_user}. {user.last_name} {user.first_name} {ExtendingUserFields.objects.get(user = user).middle_name}"))
                
                courses = UserRegisteredCourse.objects.filter(user = user).select_related('course').order_by('-timeregistered')
                
                for course in courses:
                    quiz = Quiz.objects.filter(course__id = course.course.id)
                    quiz_attempts = LearningQuizAttempt.objects.filter(user = user, quiz__in = quiz)

                    try:
                        best = (f"Дата: {localize(quiz_attempts.values_list('timefinish', flat = True).order_by('-grade').first())}\nИтоговая оценка: {int(quiz_attempts.values_list('grade', flat = True).order_by('-grade').first())} вопросов из {quiz.filter(name__contains = 'Проверка знаний').values_list('max_grade', flat = True).first()}")
                    except:
                        best = ''
                    items = (
                        ('Начало обучения:', localize(quiz_attempts.order_by('id').values_list('timestart', flat = True).first())),
                        ('Затрачено времени:', get_learning_time(course.course.id, user)),
                        ('Последний вход:', localize(last_login(user.id))),
                        ('Кол-во попыток "Проверка знаний":', quiz_attempts.count()),
                        ('Изучено вопросов:', f"{get_studied_questions(course.course.id, user)} из {question_count(course.course.id, user)}"),
                        ('Осталось дней:', (str(course.timefinish - datetime.datetime.now()).split(' ', 1)[0] + ' дней')),
                        ('Вероятность сдачи аттестации:', f"{get_percent(course.course.id, user)}%" ),
                        ('Результат лучшего теста "Проверка знаний"', best),
                    )                    
                    doc.add_paragraph()
                    
                    table = doc.add_table(1, len(items[0]))
                    table.style = 'Table Grid'
                    head_cells = table.rows[0].cells
                    for i, item in enumerate(['Курс:', course.course.name]):
                        p = head_cells[i].paragraphs[0]
                        p.add_run(item).bold = True
                        p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                    for row in items:
                        cells = table.add_row().cells
                        for i, item in enumerate(row):
                            cells[i].text = str(item)
                            if i == 2:
                                cells[i].paragraphs[0].runs[0].font.name = 'Arial'                    

            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = 'attachment; filename=reportgroup.docx'
            doc.save(response)

            return response
                
        if request.POST.get("form_type") == 'search_btn' and request.POST.get('search') != []:
            
            return HttpResponseRedirect("?search=%s" % request.POST.get('search'))

        if request.POST.get("form_type") == 'exportuser' and request.POST.getlist('checkeduser') != []:

            checkedusers = request.POST.getlist('checkeduser')
            wb = Workbook()
            ws = wb.active
            ws.append(['Имя', 'Фамилия', 'Отчество', 'Организация', 'Почта', 'Логин', 'Пароль'])
            
            for checkuser in checkedusers:
                curruser = ExtendingUserFields.objects.get(user = checkuser)
                ws.append([curruser.user.first_name, curruser.user.last_name, curruser.middle_name, curruser.department, curruser.user.email, curruser.user.username, curruser.password])
            
            response = HttpResponse(content_type='application/vnd.ms-excel')
            response['Content-Disposition'] = 'attachment;filename="userspasswords.xlsx"'
            wb.save(response)

            return response
        
        if request.POST.get("form_type") == 'senduser' and request.POST.getlist('checkeduser') != []:
            
            checkedusers = request.POST.getlist('checkeduser')
            
            for checkuser in checkedusers:
                curruser = ExtendingUserFields.objects.get(user = checkuser)
                email = EmailMessage(
                    'Доступ к сервису «Школа главного инженера»', 
                    user_message(curruser.user.username, curruser.password, curruser.user.first_name, curruser.middle_name, self.request.get_host()), 
                    to=[curruser.user.email]
                    )
                email.send()
            messages.success(request, 'Логины и пароли отправлены на эл. почту пользователя(ей)')
            return redirect('managementapp:users')
        
        if request.POST.get("form_type") == 'formimport' and 'importfile' in request.FILES:
                    
            doc = request.FILES
            myfile = doc['importfile']
            fs = FileSystemStorage()
            filename = fs.save(myfile.name, myfile)
            uploaded_file_url = []
            if filename[-4:] == 'xlsx' or filename[-3:] == 'xls':
                loadfile = load_workbook(str(BASE_DIR) + '/media/' + str(filename))
                last_group, create = Group.objects.get_or_create(
                    name = self.request.user
                )
                for sheet in loadfile.sheetnames:
                    for i in range(1,loadfile[sheet].max_row+1):
                        if loadfile[sheet].cell(row=i, column=1).value != None:

                            patternstring = re.compile("^[а-яА-ЯЁёa-zA-Z0-9]{1,15}$")
                            patternorg = re.compile("^[а-яА-ЯЁёa-zA-Z0-9- \"\'<>`.()№#«»]{1,99}$")
                            patternemail = re.compile("^[a-zA-Z0-9.!#$%&'*+/=?^_`{|}~-]+@[a-zA-Z0-9](?:[a-zA-Z0-9-]{0,61}[a-zA-Z0-9])?(?:\.[a-zA-Z0-9](?:[a-zA-Z0-9-]{0,61}[a-zA-Z0-9])?)*$")
                            
                            if not patternstring.match(loadfile[sheet].cell(row=i, column=1).value):

                                fs.delete(str(BASE_DIR) + '/media/' + str(filename))
                                messages.warning(request, 'Не верный формат загружаемых данных! Имя')
                                return redirect('managementapp:users') 

                            if not patternstring.match(loadfile[sheet].cell(row=i, column=2).value):
                                print(i, loadfile[sheet].cell(row=i, column=2).value)
                                fs.delete(str(BASE_DIR) + '/media/' + str(filename))
                                messages.warning(request, 'Не верный формат загружаемых данных! Фамилия')
                                return redirect('managementapp:users')
                            
                            if not patternstring.match(loadfile[sheet].cell(row=i, column=3).value):

                                fs.delete(str(BASE_DIR) + '/media/' + str(filename))
                                messages.warning(request, 'Не верный формат загружаемых данных! Отчество')
                                return redirect('managementapp:users')

                            if not patternorg.match(loadfile[sheet].cell(row=i, column=4).value):

                                fs.delete(str(BASE_DIR) + '/media/' + str(filename))
                                messages.warning(request, 'Не верный формат загружаемых данных! Организация')
                                return redirect('managementapp:users')

                            if not patternemail.match(loadfile[sheet].cell(row=i, column=5).value):
                                
                                fs.delete(str(BASE_DIR) + '/media/' + str(filename))
                                messages.warning(request, 'Не верный формат загружаемых данных! Почта')
                                return redirect('managementapp:users')

                            uploaded_file_url.append(loadfile[sheet].cell(row=i, column=1).value)
                            
                            if str(loadfile[sheet].cell(row=i, column=5).value).split(' ', 1)[0] == '':
                                email = str(loadfile[sheet].cell(row=i, column=5).value).split(' ', 1)[1]
                            else:
                                email = str(loadfile[sheet].cell(row=i, column=5).value)

                            password = BaseUserManager().make_random_password(length=12, allowed_chars='abcdefghjkmnpqrstuvwxyzABCDEFGHJKLMNPQRSTUVWXYZ0123456789')
            
                            last_user, create = User.objects.get_or_create(
                                username = (str(self.request.user).split('@')[0] + '_user' + str(User.objects.filter(pk__in = last_group.user_set.all()).count())),
                                first_name = loadfile[sheet].cell(row=i, column=1).value,
                                last_name = loadfile[sheet].cell(row=i, column=2).value,
                                email = email,
                                is_superuser = False,
                                is_staff = False,
                                is_active = True,
                            )
                            if create:
                                last_user.set_password(password)  
                                last_user.save()

                            last_group.user_set.add(last_user)

                            extend_user, create = ExtendingUserFields.objects.get_or_create(
                                user = last_user,
                                middle_name = loadfile[sheet].cell(row=i, column=3).value,
                                department = loadfile[sheet].cell(row=i, column=4).value,
                                password = password,
                            )
                            
                fs.delete(str(BASE_DIR) + '/media/' + str(filename))
                return redirect('managementapp:users')   

            else:

                fs.delete(str(BASE_DIR) + '/media/' + str(filename))
                messages.warning(request, 'Формат файла xlsx или xls!') 
                return redirect('managementapp:users')

        elif request.POST.get("form_type") == 'formimport' and not 'importfile' in request.FILES:

            messages.warning(request, 'Не выбран файл для импорта!')
            return redirect('managementapp:users')
        
        elif request.POST.get("form_type") != 'formimport' and request.POST.getlist('checkeduser') == []:
                
            messages.warning(request, 'Выберите пользователя!')
            return redirect('managementapp:users')
        


    def get_context_data(self, *args, **kwargs):
        context = super().get_context_data(**kwargs)

        return context

@register.filter
def check_course_end(course, user):
    return True if UserRegisteredCourse.objects.get(user = user, course = course).timefinish >= datetime.datetime.now() else False

class UserCourses(CheckManagerPermissions, ListView):
    model = UserRegisteredCourse
    template_name = 'managementapp/usercourses.html'
    paginate_by = 10
    context_object_name = 'object_list'
    
    def get_queryset(self, **kwargs):
        return (
            UserRegisteredCourse.objects
            .filter(user = self.kwargs['user_id'])
            .select_related('course')
            .order_by('-timeregistered')
            )

    def post(self, request, *args, **kwargs):
        
        if request.POST.get("form_type") == 'edituser':

            return redirect('managementapp:edituser', user_id = self.kwargs['user_id'])
                
        elif request.POST.get("form_type") == 'senduser':
            
            curruser = ExtendingUserFields.objects.get(user = self.kwargs['user_id'])
            email = EmailMessage(
                'Доступ к сервису «Школа главного инженера»', 
                user_message(curruser.user.username, curruser.password, curruser.user.first_name, curruser.middle_name, self.request.get_host()), 
                to=[curruser.user.email]
                )
            email.send()
            messages.success(request, 'Логины и пароли отправлены на эл. почту пользователя(ей)')
            return super(UserCourses, self).get(request, *args, **kwargs)

    def get_context_data(self, *args, **kwargs):
        context = super().get_context_data(**kwargs)
        context['user_id'] = self.kwargs['user_id']
        context['curr_user'] = User.objects.get(pk = self.kwargs['user_id'])
        
        return context

@register.filter
def best_percent_result(user_grade, attempt):
    grade = Quiz.objects.get(pk = LearningQuizAttempt.objects.get(pk = attempt).quiz.id)
    quotient = user_grade / grade.max_grade
    percent = quotient * 100
    return int(str(percent)[:-4])

class BestReport(CheckManagerPermissions, DetailView):
    model = LearningQuizAttempt
    context_object_name = 'object'
    template_name = 'managementapp/report.html'

    def get_object(self, **kwargs):
        return LearningQuizAttempt.objects.filter(
            user = self.kwargs['user_id'], 
            quiz = Quiz.objects.get(course_id = self.kwargs['course_id'], name__contains = 'Проверка знаний')
            ).order_by('-grade').first()
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['user_id'] = self.kwargs['user_id']
        context['quizattempt_id'] = self.get_object().id
        context['extend_user'] = ExtendingUserFields.objects.get(user = self.kwargs['user_id'])
        context['attempt_questions'] = Question.objects.filter(
            pk__in = LearningQuizAttemptGrade.objects.filter(quiz_attempt = self.get_object()).values('question')
            )
        
        return context

@user_passes_test(lambda u: u.groups.filter(name='manager').exists() == True, login_url='/')
def CourseReport(request, *args, **kwargs):
    curr_user = User.objects.get(pk = kwargs['user_id'])
    extend_user = ExtendingUserFields.objects.get(user = curr_user)
    best_attempt = LearningQuizAttempt.objects.filter(user = kwargs['user_id'], quiz = Quiz.objects.get(course_id = kwargs['course_id'], name__contains = 'Проверка знаний')).order_by('-grade').first()
    best_attempt_questions = Question.objects.filter(pk__in = LearningQuizAttemptGrade.objects.filter(quiz_attempt = best_attempt).values('question'))
    wb = Workbook()
    ws = wb.active
    img = Image(str(BASE_DIR) + '/static/managementapp/css/шгич.png')
    ws.merge_cells('B1:C1')
    ws['B1'] = str(
        'Область: ' + Course.objects.get(pk = kwargs['course_id']).full_name + '\n' + 
        'Организация: ' + extend_user.department + '\n' + 
        'ФИО: ' + curr_user.last_name + ' ' + curr_user.first_name + ' ' + extend_user.middle_name
        )
    ws.cell(row = 1, column = 2).alignment = Alignment(wrapText = True)
    ws['D1'] = (
        'Дата: ' + str(best_attempt.timestart.strftime("%d.%m.%Y %H:%M:%S")) + '\n' +
        'Итоговая оценка: ' + ('Сдан' if  best_percent_result(best_attempt.grade, Quiz.objects.get(course_id = kwargs['course_id'], name__contains = 'Проверка знаний').slug) >= 80 else 'Не сдан') + '\n' +
        '(' + str(best_attempt.grade) + ' из ' + str(Quiz.objects.get(course_id = kwargs['course_id'], name__contains = 'Проверка знаний').max_grade) + ')'
        )
    ws.cell(row = 1, column = 4).alignment = Alignment(wrapText = True)
    ws.append(['#', 'Вопрос', 'Ответ', 'Правильный ответ'])
    for idx, question in enumerate(best_attempt_questions, start = 1):
        curr_question_answers = Answer.objects.filter(pk__in = LearningQuizAttemptGrade.objects.filter(question = question, quiz_attempt = best_attempt, status = 'checked').values('answer'))
        
        checked_answer_string = ''
        for answer in curr_question_answers:
            checked_answer_string = str(checked_answer_string + (answer.text + '\n'))
        
        answer_string = ''
        for answer in question.answer_set.all():
            if answer.weight != 0:
                answer_string = str(answer_string + (answer.text + '\n'))
        
        ws.append([idx, question.text, checked_answer_string, answer_string])
        ws.cell(row = ws._current_row, column = 2).alignment = Alignment(wrapText = True)
        ws.cell(row = ws._current_row, column = 3).alignment = Alignment(wrapText = True)
        ws.cell(row = ws._current_row, column = 4).alignment = Alignment(wrapText = True)
    
    for idx, col in enumerate(ws.columns, start = 1):
        ws.column_dimensions[get_column_letter(idx)].width = 100
        for cell in col:
            alignment_obj = cell.alignment.copy(horizontal='left', vertical='center')
            cell.alignment = alignment_obj
    
    response = HttpResponse(content_type='application/vnd.ms-excel')
    response['Content-Disposition'] = 'attachment;filename="report.xlsx"'
    
    wb.save(response)

    return response    

class AddCourse(CheckManagerPermissions, ListView):
    model = Course
    template_name = 'managementapp/addcourse.html'
    context_object_name = 'object_list'

    def get_queryset(self, **kwargs):
        return Course.objects.filter(parent = None).exclude(pk__in = UserRegisteredCourse.objects.filter(user = self.kwargs['user_id']).values('course'))

    def post(self, request, *args, **kwargs):
        if request.POST.get('zadachi') == 'on':
            zadachi = True
        else:
            zadachi = False
        
        if request.POST.get('categoryid') == None and request.POST.get('zadachi') != 'on':
            last, create = UserRegisteredCourse.objects.get_or_create(
                timeregistered = datetime.datetime.now(),
                user = User.objects.get(pk = self.kwargs['user_id']),
                course = Course.objects.get(pk = request.POST.get('courseid')),
                manager = self.request.user,
                status = None,
                vot = True,
                vot_qty = 25,
            )
        elif request.POST.get('zadachi') == 'on':
            last, create = UserRegisteredCourse.objects.get_or_create(
                timeregistered = datetime.datetime.now(),
                user = User.objects.get(pk = self.kwargs['user_id']),
                course = Course.objects.get(pk = request.POST.get('courseid')),
                manager = self.request.user,
                status = None,
                problem = zadachi,
            )
        else:
            last, create = UserRegisteredCourse.objects.get_or_create(
                timeregistered = datetime.datetime.now(),
                user = User.objects.get(pk = self.kwargs['user_id']),
                course = Course.objects.get(pk = request.POST.get('categoryid')),
                manager = self.request.user,
                status = None,
                vot = True,
                vot_qty = 25,
            )

        return redirect('managementapp:usercourses', user_id = self.kwargs['user_id'])
        

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['user_id'] = self.kwargs['user_id']
        context['course_types'] = CourseType.objects.all()
        context['parent_course'] = Course.objects.all().exclude(pk__in = UserRegisteredCourse.objects.filter(user = self.kwargs['user_id']).values('course')).exclude(parent = None).order_by('name')

        return context

class ActiveCourse(CheckManagerPermissions, View):

    def get(self, request, *args, **kwargs):
        
        try:
            if UserRegisteredCourse.objects.filter(user = self.kwargs['user_id'], manager = self.request.user, course = self.kwargs['course_id'], status = None).exists():
                active_course = UserRegisteredCourse.objects.get(user = self.kwargs['user_id'], manager = self.request.user, course = self.kwargs['course_id'], status = None)
                
                if Subscription.objects.filter(manager = self.request.user).aggregate(Sum('qty'))['qty__sum'] > 0:
                    vot = 'Тренинг ВОТ'
                else:
                    messages.warning(request, 'Ошибка активации! Нет подписки!')
                    return redirect('managementapp:usercourses', user_id = self.kwargs['user_id'])
                                
                active_course.timeactive = datetime.datetime.now()
                active_course.timefinish = (datetime.datetime.now() + datetime.timedelta(days = 60))
                active_course.status = Status.objects.get(name = 'В процессе')
                
                Subscription.objects.create(
                    manager = self.request.user,
                    user = User.objects.get(pk = self.kwargs['user_id']),
                    course = Course.objects.get(pk = self.kwargs['course_id']),
                    qty = -1,
                    date = datetime.datetime.now(),
                    subtype = vot,
                )
                active_course.save()
                
                return redirect('managementapp:usercourses', user_id = self.kwargs['user_id'])

            elif UserRegisteredCourse.objects.filter(user = self.kwargs['user_id'], manager = self.request.user, course = self.kwargs['course_id'], timefinish__lte = datetime.datetime.now()).exists():
                active_course = UserRegisteredCourse.objects.get(user = self.kwargs['user_id'], manager = self.request.user, course = self.kwargs['course_id'])

                if Subscription.objects.filter(manager = self.request.user).aggregate(Sum('qty'))['qty__sum'] > 0:
                    vot = 'Тренинг ВОТ'
                else:
                    messages.warning(request, 'Ошибка активации! Нет подписки!')
                    return redirect('managementapp:usercourses', user_id = self.kwargs['user_id'])
                
                active_course.timefinish = (datetime.datetime.now() + datetime.timedelta(days = 60))
                active_course.status = Status.objects.get(name = 'В процессе')
                
                Subscription.objects.create(
                    manager = self.request.user,
                    user = User.objects.get(pk = self.kwargs['user_id']),
                    course = Course.objects.get(pk = self.kwargs['course_id']),
                    qty = -1,
                    date = datetime.datetime.now(),
                    subtype = vot,
                )
                active_course.save()
                
                return redirect('managementapp:usercourses', user_id = self.kwargs['user_id'])
            
        except:
            messages.warning(request, 'Ошибка активации!')
            return redirect('managementapp:usercourses', user_id = self.kwargs['user_id'])
    
class ActiveUser(CheckManagerPermissions, View):

    def get(self, request, *args, **kwargs):
        
        try:
            active_user = User.objects.get(pk = self.kwargs['user_id'], is_active = True)
            active_user.is_active = False
            active_user.save()

            return redirect('managementapp:users')
        except:
            active_user = User.objects.get(pk = self.kwargs['user_id'], is_active = False)
            active_user.is_active = True
            active_user.save()

            return redirect('managementapp:users')

class DeleteCourse(CheckManagerPermissions, View):

    def get(self, request, *args, **kwargs):
        
        try:
            active_course = UserRegisteredCourse.objects.get(user = self.kwargs['user_id'], manager = self.request.user, course = self.kwargs['course_id'], status = None)
            active_course.delete()


            return redirect('managementapp:usercourses', user_id = self.kwargs['user_id'])
        except:
            
            messages.warning(request, 'Ошибка удаления!')
            return redirect('managementapp:usercourses', user_id = self.kwargs['user_id'])

class CourseStatistics(CheckManagerPermissions, DetailView):
    model = Course
    template_name = 'managementapp/statistics.html'
    context_object_name = 'object'

    def get_object(self, **kwargs):
        return Course.objects.get(pk = self.kwargs['course_id'])

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['user_id'] = self.kwargs['user_id']
        context['course_id'] = self.kwargs['course_id']
        context['reg'] = UserRegisteredCourse.objects.get(user = self.kwargs['user_id'], course = self.kwargs['course_id'])
        context['time_left'] = str(UserRegisteredCourse.objects.get(user = self.kwargs['user_id'], course = self.kwargs['course_id']).timefinish - datetime.datetime.now()).split(' ', 1)[0]
        context['test_count'] = LearningQuizAttempt.objects.filter(user = self.kwargs['user_id'], quiz = Quiz.objects.get(course = self.kwargs['course_id'], name__contains = 'Проверка знаний')).count()

        return context
